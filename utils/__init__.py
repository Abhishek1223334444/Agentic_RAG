# Utility functions for the Agentic RAG Chatbot

import os
import uuid
import hashlib
from typing import List, Dict, Any, Optional
from datetime import datetime
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def generate_id() -> str:
    """Generate a unique ID."""
    logger.debug("Generating unique ID")
    id_value = str(uuid.uuid4())
    logger.debug(f"Generated ID: {id_value}")
    return id_value


def generate_hash(content: str) -> str:
    """Generate a hash for content."""
    logger.debug(f"Generating hash for content of length: {len(content)}")
    hash_value = hashlib.md5(content.encode()).hexdigest()
    logger.debug(f"Generated hash: {hash_value}")
    return hash_value


def ensure_directory(path: str) -> None:
    """Ensure a directory exists."""
    logger.debug(f"Ensuring directory exists: {path}")
    try:
        os.makedirs(path, exist_ok=True)
        logger.debug(f"Directory ensured: {path}")
    except Exception as e:
        logger.error(f"Error creating directory {path}: {e}")
        raise


def format_timestamp(timestamp: datetime) -> str:
    """Format a timestamp for display."""
    logger.debug(f"Formatting timestamp: {timestamp}")
    formatted = timestamp.strftime("%Y-%m-%d %H:%M:%S")
    logger.debug(f"Formatted timestamp: {formatted}")
    return formatted


def clean_text(text: str) -> str:
    """Clean text by removing extra whitespace and special characters."""
    logger.debug(f"Cleaning text of length: {len(text)}")
    import re
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s.,!?;:-]', '', text)
    cleaned = text.strip()
    logger.debug(f"Text cleaned: {len(text)} -> {len(cleaned)} characters")
    return cleaned


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks."""
    logger.info(f"Chunking text: length={len(text)}, chunk_size={chunk_size}, overlap={overlap}")
    
    if len(text) <= chunk_size:
        logger.debug("Text shorter than chunk size, returning as single chunk")
        return [text]
    
    chunks = []
    start = 0
    chunk_count = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at sentence boundary
        if end < len(text):
            # Look for sentence endings within the last 100 characters
            search_start = max(start + chunk_size - 100, start)
            sentence_end = text.rfind('.', search_start, end)
            if sentence_end > search_start:
                end = sentence_end + 1
                logger.debug(f"Breaking at sentence boundary: {sentence_end}")
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
            chunk_count += 1
            logger.debug(f"Created chunk {chunk_count}: {len(chunk)} characters")
        
        start = end - overlap
        if start >= len(text):
            break
    
    logger.info(f"Text chunking completed: {len(chunks)} chunks created")
    return chunks


def extract_metadata_from_filename(filename: str) -> Dict[str, Any]:
    """Extract metadata from filename."""
    logger.debug(f"Extracting metadata from filename: {filename}")
    import os
    from pathlib import Path
    
    path = Path(filename)
    metadata = {
        "filename": filename,
        "basename": path.stem,
        "extension": path.suffix,
        "size": os.path.getsize(filename) if os.path.exists(filename) else 0,
        "created_time": datetime.fromtimestamp(os.path.getctime(filename)) if os.path.exists(filename) else None
    }
    logger.debug(f"Extracted metadata: {metadata}")
    return metadata


def validate_file_type(filename: str, allowed_extensions: List[str]) -> bool:
    """Validate if file type is allowed."""
    logger.debug(f"Validating file type: {filename} against {allowed_extensions}")
    from pathlib import Path
    extension = Path(filename).suffix.lower()
    is_valid = extension in allowed_extensions
    logger.debug(f"File type validation result: {is_valid} (extension: {extension})")
    return is_valid


def get_file_size_mb(filename: str) -> float:
    """Get file size in MB."""
    logger.debug(f"Getting file size in MB: {filename}")
    if not os.path.exists(filename):
        logger.warning(f"File does not exist: {filename}")
        return 0.0
    
    size_bytes = os.path.getsize(filename)
    size_mb = size_bytes / (1024 * 1024)
    logger.debug(f"File size: {size_mb:.2f} MB ({size_bytes} bytes)")
    return size_mb


def safe_filename(filename: str) -> str:
    """Create a safe filename by removing special characters."""
    logger.debug(f"Creating safe filename from: {filename}")
    import re
    # Remove or replace special characters
    safe_name = re.sub(r'[^\w\s.-]', '_', filename)
    # Remove multiple underscores
    safe_name = re.sub(r'_+', '_', safe_name)
    result = safe_name.strip('_')
    logger.debug(f"Safe filename created: {filename} -> {result}")
    return result


class DocumentProcessor:
    """Utility class for document processing."""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file."""
        logger.info(f"Extracting text from PDF: {file_path}")
        try:
            import PyPDF2
            logger.debug("Opening PDF file")
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    logger.error(f"PDF {file_path} is password-protected")
                    raise ValueError("PDF is password-protected. Please provide an unencrypted PDF.")
                
                text = ""
                page_count = len(pdf_reader.pages)
                logger.debug(f"PDF has {page_count} pages")
                
                if page_count == 0:
                    logger.warning(f"PDF {file_path} has no pages")
                    raise ValueError("PDF has no pages or is corrupted.")
                
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            logger.debug(f"Extracted text from page {i+1}: {len(page_text)} characters")
                        else:
                            logger.warning(f"No text extracted from page {i+1} - may be image-based or empty")
                    except Exception as page_error:
                        logger.warning(f"Error extracting text from page {i+1}: {page_error}")
                        continue
                
                extracted_text = text.strip()
                
                if not extracted_text or len(extracted_text) < 10:
                    logger.error(f"PDF text extraction produced very little text ({len(extracted_text)} chars). PDF may be image-based (scanned).")
                    raise ValueError(
                        "Could not extract sufficient text from PDF. "
                        "The PDF may be image-based (scanned) or contain only images. "
                        "Please use a PDF with selectable text or use OCR to convert scanned PDFs."
                    )
                
                logger.info(f"PDF text extraction completed: {len(extracted_text)} total characters from {page_count} pages")
                return extracted_text
                
        except ValueError:
            # Re-raise ValueError as-is (these are our custom error messages)
            raise
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}", exc_info=True)
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from TXT file."""
        logger.info(f"Extracting text from TXT: {file_path}")
        try:
            logger.debug("Opening TXT file")
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read().strip()
                logger.info(f"TXT text extraction completed: {len(text)} characters")
                return text
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        logger.info(f"Extracting text from DOCX: {file_path}")
        try:
            from docx import Document
            logger.debug("Opening DOCX file")
            doc = Document(file_path)
            text = ""
            paragraph_count = len(doc.paragraphs)
            logger.debug(f"DOCX has {paragraph_count} paragraphs")
            
            for i, paragraph in enumerate(doc.paragraphs):
                paragraph_text = paragraph.text
                text += paragraph_text + "\n"
                logger.debug(f"Extracted text from paragraph {i+1}: {len(paragraph_text)} characters")
            
            logger.info(f"DOCX text extraction completed: {len(text)} total characters from {paragraph_count} paragraphs")
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    @classmethod
    def extract_text(cls, file_path: str, file_type: str) -> str:
        """Extract text from file based on type."""
        logger.info(f"Extracting text from {file_type} file: {file_path}")
        try:
            if file_type.lower() == 'pdf':
                logger.debug("Using PDF extraction method")
                return cls.extract_text_from_pdf(file_path)
            elif file_type.lower() == 'txt':
                logger.debug("Using TXT extraction method")
                return cls.extract_text_from_txt(file_path)
            elif file_type.lower() == 'docx':
                logger.debug("Using DOCX extraction method")
                return cls.extract_text_from_docx(file_path)
            else:
                logger.error(f"Unsupported file type: {file_type}")
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error in extract_text for {file_path}: {e}")
            raise


class AudioUtils:
    """Utility class for audio processing."""
    
    @staticmethod
    def convert_audio_format(audio_data: bytes, input_format: str, output_format: str) -> bytes:
        """Convert audio from one format to another."""
        logger.info(f"Converting audio format: {input_format} -> {output_format}, size={len(audio_data)} bytes")
        try:
            from pydub import AudioSegment
            import io
            
            logger.debug("Loading audio from bytes")
            # Load audio from bytes
            audio = AudioSegment.from_file(io.BytesIO(audio_data), format=input_format)
            
            logger.debug(f"Audio loaded: duration={len(audio)}ms, channels={audio.channels}, sample_rate={audio.frame_rate}")
            
            # Convert to output format
            logger.debug(f"Converting to {output_format} format")
            output_buffer = io.BytesIO()
            audio.export(output_buffer, format=output_format)
            
            result = output_buffer.getvalue()
            logger.info(f"Audio format conversion completed: {len(result)} bytes")
            return result
        except Exception as e:
            logger.error(f"Error converting audio format: {e}")
            raise
    
    @staticmethod
    def normalize_audio(audio_data: bytes, sample_rate: int = 16000) -> bytes:
        """Normalize audio to standard format."""
        logger.info(f"Normalizing audio: size={len(audio_data)} bytes, target_sample_rate={sample_rate}")
        try:
            from pydub import AudioSegment
            import io
            
            logger.debug("Loading audio for normalization")
            # Load audio
            audio = AudioSegment.from_file(io.BytesIO(audio_data))
            
            logger.debug(f"Original audio: duration={len(audio)}ms, channels={audio.channels}, sample_rate={audio.frame_rate}")
            
            # Convert to mono and set sample rate
            logger.debug("Converting to mono and setting sample rate")
            audio = audio.set_channels(1).set_frame_rate(sample_rate)
            
            logger.debug(f"Normalized audio: duration={len(audio)}ms, channels={audio.channels}, sample_rate={audio.frame_rate}")
            
            # Export as WAV
            logger.debug("Exporting as WAV")
            output_buffer = io.BytesIO()
            audio.export(output_buffer, format="wav")
            
            result = output_buffer.getvalue()
            logger.info(f"Audio normalization completed: {len(result)} bytes")
            return result
        except Exception as e:
            logger.error(f"Error normalizing audio: {e}")
            raise
